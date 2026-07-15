from __future__ import annotations

from io import BytesIO

from hypothesis_factory.models import Hypothesis


def build_docx_report(hypotheses: list[Hypothesis]) -> bytes:
    try:
        from docx import Document
    except Exception as exc:
        raise RuntimeError("python-docx is required for DOCX export") from exc

    doc = Document()
    doc.add_heading("Hypothesis Factory report", 0)
    doc.add_paragraph("Гипотезы являются кандидатами на проверку, а не доказанными научными фактами.")
    for item in hypotheses:
        doc.add_heading(f"{item.id}: {item.title}", level=1)
        doc.add_paragraph(item.hypothesis)
        doc.add_paragraph(f"Origin zone: {item.origin_uncertainty_zone.get('id')} / {item.origin_uncertainty_zone.get('type')}")
        doc.add_paragraph(f"Final score: {item.scores.get('final_score')}")
        doc.add_heading("Evidence", level=2)
        for evidence in item.evidence:
            doc.add_paragraph(f"{evidence.get('source_id')}: {evidence.get('quote')}", style="List Bullet")
        doc.add_heading("Minimal experiment", level=2)
        for step in item.minimal_experiment:
            doc.add_paragraph(step, style="List Number")
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

