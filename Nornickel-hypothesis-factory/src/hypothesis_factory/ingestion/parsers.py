from __future__ import annotations

import io
from pathlib import Path
from typing import BinaryIO

import pandas as pd

from hypothesis_factory.models import SourceDocument


def _read_bytes(file_obj: BinaryIO | io.BytesIO) -> bytes:
    file_obj.seek(0)
    return file_obj.read()


def parse_uploaded_file(file_obj: BinaryIO, filename: str) -> SourceDocument:
    suffix = Path(filename).suffix.lower()
    raw = _read_bytes(file_obj)
    title = Path(filename).name
    try:
        if suffix == ".txt":
            text = raw.decode("utf-8", errors="replace")
        elif suffix == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(raw))
            text = "\n".join(page.extract_text() or "" for page in reader.pages)
        elif suffix == ".docx":
            from docx import Document

            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
        elif suffix in {".csv", ".xlsx", ".xls"}:
            if suffix == ".csv":
                df = pd.read_csv(io.BytesIO(raw))
            else:
                df = pd.read_excel(io.BytesIO(raw))
            text = df.to_csv(index=False)
        else:
            raise ValueError(f"Unsupported file type: {suffix}")
    except Exception as exc:  # UI shows this message to the user.
        raise ValueError(f"Could not parse {title}: {exc}") from exc
    return SourceDocument(id=f"UP-{abs(hash(title)) % 100000}", title=title, text=text, source_type="upload")

